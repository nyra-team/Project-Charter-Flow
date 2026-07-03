#!/usr/bin/env python3
"""
Generate a Granules "Internal Approval Note" (NFA) .docx from a JSON payload.

Usage:
    python3 generate_nfa.py --in nfa.json --out nfa.docx

The JSON shape mirrors the pmo_nfas row:
{
  "noteNo": "01", "department": "IT", "noteDate": "06-04-2026",
  "location": "Corporate Office", "locationRequired": "HO",
  "subject": "...", "background": "...",
  "requirementItems": [{"item": "Platform Fee", "details": "..."}, ...],
  "orderFormNote": "...", "totalUsd": "$20,000", "totalInr": "₹17,00,000",
  "recommendation": "...",
  "signatories": [{"role": "Requestor", "name": "...", "status": "approved"}, ...]
}
"""
import argparse
import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x0E, 0x7C, 0x86)   # Granules teal/cyan
INK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)
GREEN = RGBColor(0x16, 0x8A, 0x4E)
RED = RGBColor(0xC0, 0x2B, 0x2B)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def run(p, text, *, bold=False, size=10, color=INK, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return r


def add_para(doc, text="", *, bold=False, size=10, color=INK, italic=False, space_after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p


def label_value(doc, label, value, *, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run(p, f"{label}: ", bold=True, size=size)
    run(p, value or "—", size=size)
    return p


def build(data, out_path):
    doc = Document()
    # Tighter margins for a form-like note
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)

    # ---- Header block (table) ----
    htbl = doc.add_table(rows=2, cols=2)
    htbl.style = "Table Grid"
    htbl.columns[0].width = Inches(4.6)
    htbl.columns[1].width = Inches(2.4)

    c = htbl.cell(0, 0)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    run(p, "Granules India Limited", bold=True, size=13, color=BRAND)
    p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    run(p2, "Internal Approval Note", bold=True, size=11, color=INK)
    p3 = c.add_paragraph()
    run(p3, "Department: ", bold=True, size=10)
    run(p3, data.get("department") or "—", size=10)

    cc = htbl.cell(0, 1)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pr = cc.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(pr, "GRANULES", bold=True, size=14, color=BRAND)

    bl = htbl.cell(1, 0)
    pb = bl.paragraphs[0]; pb.paragraph_format.space_after = Pt(2)
    run(pb, "Note No: ", bold=True, size=10); run(pb, str(data.get("noteNo") or "—"), size=10)
    pb2 = bl.add_paragraph()
    run(pb2, "Location: ", bold=True, size=10); run(pb2, data.get("location") or "—", size=10)

    br = htbl.cell(1, 1)
    pbr = br.paragraphs[0]; pbr.paragraph_format.space_after = Pt(2)
    run(pbr, "Date: ", bold=True, size=10); run(pbr, data.get("noteDate") or "—", size=10)
    pbr2 = br.add_paragraph()
    run(pbr2, "Location Required: ", bold=True, size=10); run(pbr2, data.get("locationRequired") or "—", size=10)

    add_para(doc, space_after=4)

    # ---- Subject ----
    ps = doc.add_paragraph(); ps.paragraph_format.space_after = Pt(8)
    run(ps, "Subject: ", bold=True, size=11)
    run(ps, data.get("subject") or "—", bold=True, size=11)

    # ---- Background ----
    if data.get("background"):
        add_para(doc, "Background:", bold=True, size=10, space_after=3)
        add_para(doc, data["background"], size=10, space_after=8)

    # ---- Requirement items table ----
    items = data.get("requirementItems") or []
    if items:
        add_para(doc, "Requirement / Details:", bold=True, size=10, space_after=4)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.columns[0].width = Inches(2.2)
        t.columns[1].width = Inches(4.8)
        hdr = t.rows[0].cells
        for i, h in enumerate(("Item", "Details")):
            set_cell_bg(hdr[i], "0E7C86")
            hp = hdr[i].paragraphs[0]
            run(hp, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        for it in items:
            row = t.add_row().cells
            run(row[0].paragraphs[0], str(it.get("item") or ""), bold=True, size=10)
            run(row[1].paragraphs[0], str(it.get("details") or ""), size=10)
        add_para(doc, space_after=6)

    # ---- Order form / attachment note ----
    if data.get("orderFormNote"):
        add_para(doc, "Order Form Details:", bold=True, size=10, space_after=3)
        add_para(doc, data["orderFormNote"], size=10, space_after=8)

    # ---- Total costing ----
    usd = (data.get("totalUsd") or "").strip()
    inr = (data.get("totalInr") or "").strip()
    if usd or inr:
        parts = [x for x in (usd, inr) if x]
        pt = doc.add_paragraph(); pt.paragraph_format.space_after = Pt(8)
        run(pt, "Total Costing = ", bold=True, size=11)
        run(pt, "  /  ".join(parts), bold=True, size=11, color=BRAND)

    # ---- Recommendation ----
    if data.get("recommendation"):
        add_para(doc, "Recommendation:", bold=True, size=10, space_after=3)
        add_para(doc, data["recommendation"], size=10, space_after=10)

    # ---- Signatory grid ----
    sigs = data.get("signatories") or []
    if sigs:
        add_para(doc, "Approvals:", bold=True, size=10, space_after=4)
        cols = 3
        rows = (len(sigs) + cols - 1) // cols
        st = doc.add_table(rows=rows, cols=cols)
        st.style = "Table Grid"
        for idx, sig in enumerate(sigs):
            cell = st.cell(idx // cols, idx % cols)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            pr = cell.paragraphs[0]; pr.paragraph_format.space_after = Pt(2)
            run(pr, str(sig.get("role") or ""), bold=True, size=9, color=GREY)
            pn = cell.add_paragraph(); pn.paragraph_format.space_after = Pt(4)
            run(pn, str(sig.get("name") or "—"), bold=True, size=10)
            # No printed status word — this document is a frozen snapshot sent for
            # e-signature; a baked "Pending"/"Approved" would go stale the moment
            # someone signs (the signature stamp itself is the approval evidence,
            # and live status lives in the app). Just a "Signature:" label + the
            # invisible [[SIGn]] anchor the Documenso sender pins the field to.
            plabel = cell.add_paragraph()
            run(plabel, "Signature:", size=8, color=GREY)
            # Invisible e-sign anchor (white text — present in the PDF text layer,
            # unseen on paper). space_after reserves room for the signature stamp.
            psig = cell.add_paragraph()
            psig.paragraph_format.space_after = Pt(30)
            run(psig, f"[[SIG{idx + 1}]]", size=8, color=RGBColor(0xFF, 0xFF, 0xFF))
        # blank trailing cells stay empty

    add_para(doc, space_after=2)
    foot = add_para(
        doc, "Generated by Granules Project Hub (PMO).",
        size=8, color=GREY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="inp", required=True)
    ap.add_argument("--out", dest="out", required=True)
    args = ap.parse_args()
    with open(args.inp, "r", encoding="utf-8") as f:
        data = json.load(f)
    build(data, args.out)
    print(args.out)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:  # noqa
        print(f"generate_nfa error: {e}", file=sys.stderr)
        sys.exit(1)
