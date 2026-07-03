#!/usr/bin/env python3
"""
Generate a consolidated Granules Project Charter + Note for Approval (NFA)
.docx from a JSON payload (the merged pmo_charters row).

Usage:
    python3 generate_charter_nfa.py --in charter.json --out charter.docx

Mirrors the 17-section MES Charter+NFA template:
  1. Project Information (sponsor / PMs / dates / category / entity / note no)
  2. Executive Summary
  3. Project Description
  4. Background — current state + business drivers
  5. Scope (In / Out)
  6. Benefits & KPIs
  7. Business Case & ROI / Annum
  8. Implementation Roadmap & Milestones
  9. Revised Project Investment Summary (one-time CAPEX/OPEX + FY-wise recurring)
 10. Earlier vs Revised NFA (previous_nfa_amount vs total / le_amount)
 11. Constraints
 12. Risks (from pmo_risks — passed in as `risks` array)
 13. Assumptions
 14. Potential Additional Budget Areas
 15. Project Governance (Steering Committee + Key Project Members)
 16. Attachments
 17. Approval & Sign-off (from charter.signatories — dynamic chain from the DOA matrix)

The sign-off block reflects the resolved DOA chain with per-row status / decided-at.
"""
import argparse
import json
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x0E, 0x7C, 0x86)   # Granules teal
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x47, 0x55, 0x69)
HDR_BG = "0E7C86"
SUBHDR_BG = "CCE7E9"
BAND_BG = "F1F5F9"
ACCENT_BG = "DBF2F4"


# ───── helpers ──────────────────────────────────────────────────────────────
_BARE_QUARTER_RE = re.compile(r"^\s*(?:Q|Quarter\s*)([1-4])\s*$", re.IGNORECASE)


def _fiscal_year_label(base_date_str):
    """Indian FY (Apr–Mar): a date in Apr-2025…Mar-2026 → 'FY26'. None if unparseable."""
    if not base_date_str:
        return None
    s = str(base_date_str).strip()
    dt = None
    for cut in (s[:10], s[:19]):
        try:
            dt = datetime.fromisoformat(cut)
            break
        except ValueError:
            continue
    if dt is None:
        return None
    fy_end = dt.year + 1 if dt.month >= 4 else dt.year
    return f"FY{fy_end % 100:02d}"


def qualify_target_date(raw, base_date_str):
    """A bare quarter ('Q2') is ambiguous in the doc — append the charter's FY so
    the timeline always names a year. Non-quarter values pass through untouched.
    ponytail: assumes all bare quarters fall in the charter's base FY; doesn't roll
    multi-FY projects (upgrade: honour a per-milestone year when the data carries one)."""
    if raw is None:
        return ""
    s = str(raw).strip()
    m = _BARE_QUARTER_RE.match(s)
    if not m:
        return s
    fy = _fiscal_year_label(base_date_str)
    return f"Q{m.group(1)} {fy}" if fy else s


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, sz=4, color="94A3B8"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        borders.append(b)
    tcPr.append(borders)


def set_cell_margins(cell, *, top=40, bottom=40, left=110, right=110):
    """Padding inside a table cell (values in twips, 20 = 1pt). Gives the
    tables proper breathing room instead of text hugging the borders."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def set_tracking(r, twips):
    """Letter-spacing on a run (twips; 20 = 1pt). Used for the airy small-caps
    section headings."""
    rPr = r._r.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), str(twips))
    rPr.append(spc)


def set_run(r, *, size=10, bold=False, color=INK, italic=False):
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20 if level == 1 else 13)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 4)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = True  # never strand a heading at a page foot
    # Level-1 = an airy letter-spaced small-caps section header (corporate look);
    # deeper levels stay sentence-case.
    r = p.add_run(text.upper() if level == 1 else text)
    set_run(r, size=(12.5 if level == 1 else 11.5 if level == 2 else 10.5),
            bold=True, color=BRAND if level <= 2 else INK)
    if level == 1:
        set_tracking(r, 24)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "4"); bot.set(qn("w:color"), "0E7C86")
        pBdr.append(bot); pPr.append(pBdr)
    return p


def para(doc, text="", *, size=10.5, bold=False, italic=False, color=INK, align=None, space_after=7):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15  # readable body leading
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def bullet(doc, text, *, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run(r, size=size)
    return p


def build_table(doc, rows, *, col_widths_cm=None, header_rows=1, first_col_header=False, band=False):
    if not rows:
        return None
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    for r_idx, row in enumerate(rows):
        is_header = r_idx < header_rows
        for c_idx, txt in enumerate(row):
            cell = t.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run("" if txt is None else str(txt))
            if is_header:
                set_run(r, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                set_tracking(r, 12)
                set_cell_bg(cell, HDR_BG)
            elif first_col_header and c_idx == 0:
                set_run(r, size=10, bold=True, color=INK)
                set_cell_bg(cell, SUBHDR_BG)
            else:
                set_run(r, size=10)
                if band and r_idx % 2 == 0:
                    set_cell_bg(cell, BAND_BG)
    return t


def fmt_inr(v):
    """1.5e7 → '₹1.50 Cr', 8e5 → '₹8.00 L', 50000 → '₹50,000'."""
    if v is None or v == "" or v is False:
        return "—"
    try:
        n = float(v)
    except (ValueError, TypeError):
        return str(v)
    if n >= 1e7:
        return f"₹{n/1e7:.2f} Cr"
    if n >= 1e5:
        return f"₹{n/1e5:.2f} L"
    return f"₹{n:,.0f}"


def fmt_date(s):
    if not s:
        return "—"
    try:
        return datetime.fromisoformat(str(s).replace("Z", "+00:00")).strftime("%d-%b-%Y")
    except Exception:
        return str(s)


# ───── build ────────────────────────────────────────────────────────────────
def build(data, out_path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.6); s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # ── Cover ──
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("GRANULES INDIA LIMITED"); set_run(r, size=10.5, bold=True, color=MUTED); set_tracking(r, 60)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Project Charter & Note for Approval (NFA)"); set_run(r, size=19, bold=True, color=BRAND); set_tracking(r, 4)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(data.get("title") or "[ Project name ]"), size=14, bold=True, color=INK)

    pcId = next((t.split(":", 1)[1] for t in (data.get("strategicAlignmentTags") or []) if isinstance(t, str) and t.startswith("PC_ID:")), None)
    rev_label = f"Revision {data.get('revision', 1)}"
    sub = f"{pcId or '—'}  ·  {rev_label}  ·  Generated {datetime.utcnow().strftime('%d-%b-%Y')}"
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_run(p.add_run(sub), size=10, italic=True, color=MUTED)

    # Sections follow the Granules ANNEXURE-I Project Charter / NFA order exactly.

    # 1. Project Information
    heading(doc, "1. Project Information", level=2)
    pm_type = (data.get("pmType") or "").strip()
    pm_name = (data.get("pmName") or "").strip()
    pm_disp = (f"{pm_type} — {pm_name}".strip(" —")) if (pm_type or pm_name) else "—"
    info_rows = [
        ["Project Name", data.get("title") or "—", "Project Sponsor", data.get("projectSponsor") or "—"],
        ["Function / Department", data.get("department") or "—", "Project Category", data.get("category") or "—"],
        ["PM", pm_disp, "Note No", data.get("noteNo") or "—"],
        ["Date of Project Approval", fmt_date(data.get("projectApprovalDate")), "Last Revision Date", fmt_date(data.get("lastRevisionDate"))],
    ]
    build_table(doc, info_rows, col_widths_cm=[4.0, 4.6, 4.0, 5.0], header_rows=0, first_col_header=True)
    heading(doc, "Key Project Members", level=2)
    members = data.get("keyProjectMembers") or []
    members_disp = ", ".join(m.get("name", "") for m in members if m.get("name")) or "—"
    para(doc, members_disp)

    # ── Reorderable body sections (2…N) ──────────────────────────────────────
    # Order is author-controlled (drag-to-reorder on the e-NFA form), persisted
    # as data["sectionOrder"]. Section 1 (Project Information) and the Approval
    # block stay fixed. Heading numbers follow the final order, not hard-coded.
    def sec_exec(n):
        heading(doc, f"{n}. Project Description / Executive Summary", level=1)
        para(doc, data.get("executiveSummary") or data.get("description") or "—")
        heading(doc, "Background", level=2)
        para(doc, data.get("background") or "—")
        return True

    def sec_scope(n):
        heading(doc, f"{n}. Scope", level=1)
        heading(doc, f"{n}.1 In Scope", level=2)
        para(doc, data.get("scope") or "—")
        heading(doc, f"{n}.2 Out of Scope", level=2)
        para(doc, data.get("outOfScope") or "—")
        return True

    def sec_outcome(n):
        heading(doc, f"{n}. Business Outcome", level=1)
        para(doc, data.get("businessOutcome") or "—")
        return True

    def sec_constraints(n):
        heading(doc, f"{n}. Constraints", level=1)
        if data.get("constraints"):
            para(doc, data["constraints"])
        cons_rows = [["Item", "Detail"]]
        if data.get("tentativeBudget"):
            cons_rows.append(["Approved Budget", fmt_inr(data.get("tentativeBudget"))])
        if data.get("leAmount"):
            cons_rows.append(["LE Budget", fmt_inr(data.get("leAmount"))])
        if data.get("scopeLimitations"):
            cons_rows.append(["Scope Limitations", data.get("scopeLimitations")])
        if len(cons_rows) > 1:
            build_table(doc, cons_rows, col_widths_cm=[5.0, 12.6], first_col_header=True)
        return True

    def sec_deliverables(n):
        heading(doc, f"{n}. Project Deliverables (Key Milestones)", level=1)
        milestones = data.get("milestones") or []
        if milestones:
            base = data.get("startDate") or data.get("projectApprovalDate") or data.get("noteDate") or data.get("createdAt")
            rows = [["Key Milestone", "Responsible", "Target Date"]]
            for m in milestones:
                rows.append([m.get("milestone", ""), m.get("responsible", ""), qualify_target_date(m.get("targetDate", ""), base)])
            build_table(doc, rows, col_widths_cm=[9.4, 4.6, 3.6], band=True)
        else:
            para(doc, "No milestones captured.", italic=True, color=MUTED)
        return True

    def sec_benefits(n):
        heading(doc, f"{n}. Benefits", level=1)
        para(doc, "Topline improvement, bottom-line optimization, compliance benefits & productivity improvement.",
             italic=True, color=MUTED, space_after=4)
        kpis = data.get("kpis") or []
        if kpis:
            rows = [["KPI", "Baseline", "Goal"]]
            for k in kpis:
                rows.append([k.get("kpi", ""), k.get("baseline", ""), k.get("goal", "")])
            build_table(doc, rows, col_widths_cm=[8.4, 4.4, 4.8], band=True)
        else:
            para(doc, "No KPIs captured.", italic=True, color=MUTED)
        heading(doc, "ROI / Annum", level=2)
        roi = data.get("roiPerAnnum")
        if roi:
            para(doc, fmt_inr(roi))
        else:
            para(doc, "Not quantified.", italic=True, color=MUTED)
        return True

    def sec_risks(n):
        heading(doc, f"{n}. Risks", level=1)
        risks = data.get("structuredRisks") or []
        risks_text = data.get("risks")
        if risks:
            for r in risks:
                bullet(doc, f"{r.get('title','')}: {r.get('description','') or ''}".strip(": "))
        elif isinstance(risks_text, str) and risks_text.strip():
            para(doc, risks_text)
        else:
            para(doc, "—")
        heading(doc, "Assumptions", level=2)
        para(doc, data.get("assumptions") or "—")
        heading(doc, "Potential Additional Budget Areas", level=2)
        para(doc, data.get("potentialAdditionalBudget") or "—")
        return True

    def sec_vendor(n):
        heading(doc, f"{n}. Vendor Comparison Matrix", level=1)
        vm = data.get("vendorMatrix") or {}
        vcols = vm.get("columns") or []
        vrows = [r for r in (vm.get("rows") or [])
                 if any((str(c).strip() if c is not None else "") for c in r)]
        if vcols and vrows:
            table = [list(vcols)] + [[("" if c is None else str(c)) for c in r] for r in vrows]
            build_table(doc, table, band=True)
        else:
            para(doc, "No vendor comparison captured.", italic=True, color=MUTED)
        return True

    def sec_additional(n):
        custom = [f for f in (data.get("customFields") or []) if (f.get("label") or f.get("value"))]
        if not custom:
            return False  # nothing to show — skip without consuming a number
        heading(doc, f"{n}. Additional Information", level=1)
        for f in custom:
            if f.get("label"):
                para(doc, f.get("label", ""), bold=True, space_after=1)
            para(doc, f.get("value", ""), space_after=4)
        return True

    RENDERERS = {
        "executiveSummary": sec_exec,
        "scope": sec_scope,
        "businessOutcome": sec_outcome,
        "constraints": sec_constraints,
        "deliverables": sec_deliverables,
        "benefits": sec_benefits,
        "risks": sec_risks,
        "vendorMatrix": sec_vendor,
        "additionalFields": sec_additional,
    }
    DEFAULT_ORDER = ["executiveSummary", "scope", "businessOutcome", "constraints",
                     "deliverables", "benefits", "risks", "vendorMatrix", "additionalFields"]
    # Honor the author's order; append any known section they omitted so the doc
    # never silently drops content. Unknown ids (e.g. "rfp", no doc section) skip.
    raw_order = data.get("sectionOrder") or []
    order = [s for s in raw_order if s in RENDERERS]
    for s in DEFAULT_ORDER:
        if s not in order:
            order.append(s)

    n = 2
    for sid in order:
        if RENDERERS[sid](n):
            n += 1

    # Approval & Sign-off
    heading(doc, "Approval & Sign-off", level=1)
    para(doc, "Signatories listed below were resolved via the active DOA matrix at the time of submission.",
         italic=True, color=MUTED, space_after=6)

    sigs = data.get("signatories") or []
    if sigs:
        # No Status / Decided At columns — this document is a frozen snapshot sent
        # for e-signature; those would bake in "Pending"/blank and go stale once
        # signing starts. The signature stamp itself is the approval evidence;
        # live status lives in the app.
        rows = [["Role", "Name", "Signature"]]
        for i, s in enumerate(sigs, start=1):
            rows.append([
                (s.get("role") or "").replace("_", " ").title(),
                s.get("name", "") or "—",
                f"[[SIG{i}]]",
            ])
        t = build_table(doc, rows, col_widths_cm=[4.0, 5.4, 6.0], band=True)
        # Invisible e-sign anchors: the Documenso sender searches the PDF text
        # layer for "[[SIGn]]" and pins signer n's signature field exactly on
        # its cell. Whiten the markers + reserve stamp height per row.
        for row in t.rows[1:]:
            cell = row.cells[-1]
            for p in cell.paragraphs:
                for r in p.runs:
                    if re.fullmatch(r"\[\[SIG\d+\]\]", r.text or ""):
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.size = Pt(8)
            pad = cell.paragraphs[-1]
            pad.paragraph_format.space_after = Pt(30)
    else:
        para(doc, "Sign-off block populated on submission — empty in draft mode.", italic=True, color=MUTED)

    # Footer
    para(doc, "", space_after=4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Project Charter & Note for Approval (NFA). Source: pmo_charters (Granules Project Hub)."),
            size=8, italic=True, color=MUTED)

    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="inp", required=True)
    ap.add_argument("--out", dest="out", required=True)
    args = ap.parse_args()
    with open(args.inp, "r", encoding="utf-8") as f:
        data = json.load(f)
    build(data, args.out)
    print(args.out)


if __name__ == "__main__":
    sys.exit(main())
